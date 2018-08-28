#!/usr/bin/env python
# coding=utf-8

__author__ = 'guankx'

__date__ = '2018-06-18'

import time

#import DBUtils

#import re

import sys

import docx

import chardet

import os
from  docx.oxml.ns import  qn

reload(sys)

#sys.setdefaultencoding('utf-8')

curr_dir = os.path.dirname(os.path.abspath(__file__))
filepath =  os.path.join(curr_dir,'files')
filepath2 = os.path.join(curr_dir,'files2')
filtered_words_txt_path = os.path.join(curr_dir,'B.txt')

class Node(object):

   def __init__(self):

       self.children = None

       #self.badword = None
       self.flag = False

       #self.isEnd = None
def add_word(root,word):

   node = root

   for i in range(len(word)):

       if node.children == None:

           node.children = {}
           node.children[word[i]] = Node()
          
      
       elif word[i] not in node.children:

           node.children[word[i]] = Node()

       node = node.children[word[i]] 
   #node.badword = word
   node.flag = True
   #node.isEnd = 1

def init():

   root = Node()

   #result = u"卧槽\n尼玛\n"

   '''#-------------------------------------------

   #从数据库中读取

   db = DBUtils.DBUtils('localhost','root','4521','test')

   db.set_table("base_badwords")

   result = db.select(['words'])

   for line in result:

       #只匹配中文/英文/数字

      #li = ''.join(re.findall(re.compile(u'[a-zA-Z0-9\u4e00-\u9fa5]'),line[0]))

       #if li:

       #    add_word(root,li.lower())

       add_word(root,line[0].lower())

   return root

   '''#-------------------------------------------

   #-------------------------------------------

    #从文件中读取

   with open ('/Users/sunflower/Downloads/B.txt','r') as result:

       for line in result:

           #只匹配中文/英文/数字

          #li = ''.join(re.findall(re.compile(u'[a-zA-Z0-9\u4e00-\u9fa5]'),line.strip().decode('utf8')))

           #if li:

           #    print li

           #    add_word(root,li.lower())

           if line.strip():
              
              add_word(root,line.strip().decode('utf8').lower())

   return root

   #'''#-------------------------------------------

def is_contain(message, root):
   res = set()
   for i in range(len(message)):

       p = root

       j = i

       
       while (j<len(message) and p.children!=None and message[j] in p.children):
           if p.flag == True:
              res.add(message[i:j])
             
              message.replace(message[i:j],'('+message[i:j]+')')
           p = p.children[message[j]]
           j = j+1

       if p.children==None:
            res.add(message[i:j])  
            message.replace(message[i:j],'('+message[i:j]+')')
       #if p.badword == message[i:j]:

           #print '--word--',p.badword,'-->',message

           #return p.badword

       #if p.isEnd:

           #return message[i:j]

   #return 0
   return res

def dfa():

   print '------------dfa start-----------'

   print 'init ...'

   root = init()

   print 'init done!'

   #message = u'卧槽'

   ''' db = DBUtils.DBUtils('localhost','root','4521','test')

   db.set_table("user_profile")#用户表

   result = db.select(['nickname','user_id'])#取 昵称 和 用户ID

   print "user count:",len(result) '''

    #开始计时

   start_time = time.time()

   data = []

   '''for line in result:

      message = ''.join(re.findall(re.compile(u'[a-zA-Z0-9\u4e00-\u9fa5]'),line[0]))

       #print '***message***',len(message)

       res = is_contain(message.lower(),root)

       if res:

           data.append([line[1],res,message])
   '''
   '''file_object = open('/Users/sunflower/Downloads/test.txt')
   file_context = file_object.read()
   message = ''.join(re.findall(re.compile(u'[a-zA-Z0-9\u4e00-\u9fa5]'),file_context.decode('utf-8')))
   #message = file_context.decode('utf-8')
       #print '***message***',len(message)
   curr_dir = os.path.dirname(os.path.abspath(__file__))
   file=docx.Document(os.path.join(curr_dir,'宠妾作死日常 文案.docx'))
   message = ''

   for i in range(len(file.paragraphs)):
        message += file.paragraphs[i].text+"\r\n"
   '''
   files = os.listdir(filepath)
   for fi in files:
      content_path = os.path.join(filepath,fi)
      content_path2 = os.path.join(filepath2,fi)
      message = ''
      outputfile = ''
      if fi.endswith('docx'):
        file=docx.Document(content_path)
     

        for i in range(len(file.paragraphs)):
            message += file.paragraphs[i].text+"\r\n"
     
        x = is_contain(message.lower(),root)
        for item in x:
           message = message.replace(item,'('+item+')')


        sys.setdefaultencoding('gb18030')
        #output = sys.stdout
        #outputfile=open(content_path2,'w')
        document = docx.Document()
        #sys.stdout=outputfile
        paragraph = document.add_paragraph(message)
        
        #设置中文字体
        run = paragraph.add_run(u'设置中文字体，')
        run.font.name=u'微软雅黑'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        document.save(content_path2);
        #print message
        #outputfile.close()
        #sys.stdout=output
      if fi.endswith('txt'):
        sys.setdefaultencoding('gb18030')
        file_object = open(content_path)
       
        try:
            message = file_object.read().encode('utf-8').decode('utf-8')
            #print type(message)
            #message = file_object.read().decode('utf-8')
            #print message
            x = is_contain(message.lower(),root)
            for item in x:
              message = message.replace(item,'('+item+')')

            #sys.setdefaultencoding('latin1')
            sys.setdefaultencoding('utf-8')
            output=sys.stdout
            outputfile=open(content_path2,'w')
            sys.stdout=outputfile
            #print message.decode('latin1')
            print message
            outputfile.close()
            sys.stdout=output
        finally:
             file_object.close()
    
      
   #if res:

   #     data.append(file_context,res,message)

   end_time = time.time()

   #这里把含有敏感词的用户存入数据库

   '''db.set_table('bad_user')#含有敏感词的用户表，其实可以直接在用户表中添加相应字段，并标记为0:正常用户，1:含有敏感词用户

   fields = ['user_id','bad_word','nickname']#用户ID 敏感词 昵称

   db.insertmany(data,fields)'''

    #输出所用时间

   print (end_time - start_time) 

if __name__ == '__main__':

    dfa()